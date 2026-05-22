import os
import shutil
import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image, Table, TableStyle, PageBreak, KeepTogether
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.pdfgen import canvas

# --------------------------------------------------------------------------
# Paths and Configuration
# --------------------------------------------------------------------------
PROJECT_DIR = os.path.dirname(os.path.abspath(__file__))
ASSETS_DIR = os.path.join(PROJECT_DIR, "assets")
DOCX_OUT = os.path.join(PROJECT_DIR, "UniGuide_Assignment_Submission.docx")
PDF_OUT = os.path.join(PROJECT_DIR, "UniGuide_Assignment_Submission.pdf")

# Copy generated mockups to assets directory if needed
src_images = {
    "welcome_screen.png": r"C:\Users\Asus\.gemini\antigravity\brain\4ca5bfe1-aed6-4909-84af-b4137fd37c3d\welcome_screen_1779344563810.png",
    "navigation_flow.png": r"C:\Users\Asus\.gemini\antigravity\brain\4ca5bfe1-aed6-4909-84af-b4137fd37c3d\navigation_flow_1779344585310.png",
    "admission_faq.png": r"C:\Users\Asus\.gemini\antigravity\brain\4ca5bfe1-aed6-4909-84af-b4137fd37c3d\admission_faq_1779344606967.png",
    "fallback_handling.png": r"C:\Users\Asus\.gemini\antigravity\brain\4ca5bfe1-aed6-4909-84af-b4137fd37c3d\fallback_handling_1779344631073.png"
}

images_ok = {}
for filename, src_path in src_images.items():
    dest_path = os.path.join(ASSETS_DIR, filename)
    if os.path.exists(src_path):
        try:
            shutil.copy2(src_path, dest_path)
            images_ok[filename] = dest_path
        except Exception as e:
            images_ok[filename] = src_path
    else:
        images_ok[filename] = dest_path if os.path.exists(dest_path) else None

title_text = "Assignment 1: Chatbot Development"
subtitle_text = "UniGuide: Intelligent College Campus Chatbot with Python Flask & Gemini AI"
author_text = "Prepared for Assignment Submission"
date_text = datetime.datetime.now().strftime("%B %d, %Y")

# --------------------------------------------------------------------------
# Word Document (.docx) Generation
# --------------------------------------------------------------------------
def generate_docx():
    print("Generating Word Document...")
    doc = Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Set styles
    styles = doc.styles
    normal_style = styles['Normal']
    normal_style.font.name = 'Arial'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(0x33, 0x41, 0x55) # Slate 700

    # Title Page
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = title_p.add_run(f"\n\n\n\n{title_text}\n")
    run_t.font.name = 'Arial'
    run_t.font.size = Pt(28)
    run_t.font.bold = True
    run_t.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5) # Indigo 600

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_s = sub_p.add_run(f"{subtitle_text}\n\n\n\n\n\n")
    run_s.font.name = 'Arial'
    run_s.font.size = Pt(16)
    run_s.font.color.rgb = RGBColor(0x64, 0x74, 0x8B) # Slate 500

    info_p = doc.add_paragraph()
    info_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_info = info_p.add_run(f"Author: Student\nDate: {date_text}\nArchitecture: Python Flask Backend & Gemini AI integration\n")
    run_info.font.name = 'Arial'
    run_info.font.size = Pt(12)
    run_info.font.bold = True
    
    doc.add_page_break()

    # Section 1
    h1 = doc.add_heading("1. Executive Summary & Project Objective", level=1)
    h1.runs[0].font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
    
    doc.add_paragraph(
        "Chatbots have emerged as an essential interface in modern web systems, facilitating instant support "
        "and reducing administrative workload. This project, titled UniGuide, is a smart campus assistant chatbot "
        "built with a Python Flask server backend and optional Google Gemini AI integration. The chatbot is designed to guide "
        "students, prospective applicants, and parents through admissions procedures, tuition structures, academic "
        "calendars, campus geography, and support tickets."
    )
    doc.add_paragraph("Key objectives of the UniGuide chatbot include:")
    doc.add_paragraph("• 24/7 Information Access: Instantly answering common student questions regarding tuition fees and term dates.", style='List Bullet')
    doc.add_paragraph("• Dual AI Processing: Leveraging Google Gemini LLM API for natural language understanding while using a robust offline rule-based regex fallback.", style='List Bullet')
    doc.add_paragraph("• Admissions Streamlining: Helping new applicants understand application steps and deadlines.", style='List Bullet')
    doc.add_paragraph("• Live Support Integration: Enabling users to raise troubleshooting tickets, gathering email data, and generating confirmation IDs via stateful sessions.", style='List Bullet')

    # Section 2
    h2 = doc.add_heading("2. Conversation Flow Design", level=1)
    h2.runs[0].font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
    doc.add_paragraph(
        "To deliver a robust conversational user experience, UniGuide uses a rule-based state-machine design. "
        "The logic splits interactions into predefined, structured flows: Welcome Greeting, Admission Flow, Navigation Flow, "
        "Ticketing Form Flow, and Fallback/Error state."
    )
    
    doc.add_heading("Detailed Conversation Node Structure", level=2)
    doc.add_paragraph("1. Welcome Node: Greets the user, details the bot's abilities, and suggests quick-reply buttons (Admissions, Fees, Map, Exams, Raise Ticket).")
    doc.add_paragraph("2. Admissions Node: Triggers on keywords like 'apply', 'admission', 'enrol'. Displays steps, deadlines, and links to the Application portal.")
    doc.add_paragraph("3. Navigation Node: Triggers on map-related keywords. Displays a simulated graphical map and popular landmarks.")
    doc.add_paragraph("4. Support Ticket State: Collects student email, validates formatting, prompts for issue details, submits to HELP-DESK queue, and returns confirmation IDs.")
    doc.add_paragraph("5. Fallback Error Node: Triggers on unmatched queries. Suggests search alternatives and displays hotkeys to return to known topics.")

    # Section 3
    h3 = doc.add_heading("3. Technical Implementation & UI Design", level=1)
    h3.runs[0].font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
    doc.add_paragraph(
        "UniGuide is built as a highly responsive Python Flask application to guarantee server-side security, "
        "live API connectivity, and a premium look-and-feel. It uses native web APIs and complies with modern aesthetic guidelines."
    )
    doc.add_paragraph("• Backend Architecture (Python Flask): Configures session state memory to store support ticket fields. Exposes routing for landing pages ('/') and API chat processing endpoints ('/api/chat'). Loads the Google Generative AI (Gemini) SDK for intelligent reasoning.", style='List Bullet')
    doc.add_paragraph("• Structure (HTML5/Templates): Rendered dynamically by Flask. Organizes the application into a split layout, showing stats cards in the sidebar and chat messages on the main viewport.", style='List Bullet')
    doc.add_paragraph("• Design & Style (Vanilla CSS3): Uses glassmorphism (frosted borders, backdrop filters) to blend components nicely in dark and light modes. Incorporates smooth transit keyframes for message bubble entries.", style='List Bullet')
    doc.add_paragraph("• Client Script (Vanilla JS): Manages frontend elements, triggers typing status indicators, interacts with the backend `/api/chat` API using Fetch, and plays sound effects.", style='List Bullet')

    doc.add_page_break()

    # Section 4
    h4 = doc.add_heading("4. Chatbot Screens & Output Analysis", level=1)
    h4.runs[0].font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
    
    # Adding screenshots
    images_to_add = [
        ("welcome_screen.png", "Figure 1: UniGuide Welcome Interface & Common Topics Panel"),
        ("admission_faq.png", "Figure 2: Admissions Requirements & Fees Inquiry Interface"),
        ("navigation_flow.png", "Figure 3: Interactive Campus Maps and Directions"),
        ("fallback_handling.png", "Figure 4: Robust Fallback Error Handling for Unrecognized Queries")
    ]
    
    for filename, caption in images_to_add:
        img_path = images_ok.get(filename)
        if img_path and os.path.exists(img_path):
            doc.add_paragraph().add_run().add_picture(img_path, width=Inches(5.0))
            caption_p = doc.add_paragraph()
            caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_cap = caption_p.add_run(caption)
            run_cap.italic = True
            run_cap.font.size = Pt(10)
            doc.add_paragraph() # Spacer
        else:
            doc.add_paragraph(f"[Image Missing: {filename} - Check setup]")

    # Project Links
    h4_links = doc.add_heading("Project Working Directory & Links", level=2)
    p_link = doc.add_paragraph()
    p_link.add_run("• Local Source Directory: ").bold = True
    project_dir_url = PROJECT_DIR.replace('\\', '/')
    p_link.add_run(f"file:///{project_dir_url}\n")
    p_link.add_run("• Flask Entry Script: ").bold = True
    flask_script_url = os.path.join(PROJECT_DIR, 'app.py').replace('\\', '/')
    p_link.add_run(f"file:///{flask_script_url}\n")
    p_link.add_run("• Local Dev Server Port: ").bold = True
    p_link.add_run("http://localhost:5000\n")

    # Table of Dialog Flows
    h4_table = doc.add_heading("Conversation Interaction Mapping", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Shading Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'User Intent / Input'
    hdr_cells[1].text = 'Pattern Matching Rule'
    hdr_cells[2].text = 'Bot Action & Output'
    
    # Make header bold
    for cell in hdr_cells:
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True

    interactions = [
        ("admission requirements", "Keyword matches: 'admission', 'apply'", "Displays application portal link, deadlines, and criteria list."),
        ("tuition fees", "Keyword matches: 'fees', 'cost', 'tuition'", "Displays cost comparison list (in-state vs out-of-state) and scholarship links."),
        ("campus map library", "Keyword matches: 'map', 'directions', 'library'", "Draws simulated campus map box with dynamic locator pulsing."),
        ("raise ticket", "Keyword matches: 'ticket', 'support', 'help'", "Initiates stateful ticket form. Prompts for student email address, then prompt for issue description, returning ID."),
        ("fly to Mars", "No matches (Fallback triggers)", "Returns friendly error notification card containing popular recommended buttons.")
    ]

    for intent, rule, out in interactions:
        row_cells = table.add_row().cells
        row_cells[0].text = intent
        row_cells[1].text = rule
        row_cells[2].text = out

    doc.add_paragraph("\n")

    # Section 5
    h5 = doc.add_heading("5. Conclusion & Future Scopes", level=1)
    h5.runs[0].font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)
    doc.add_paragraph(
        "UniGuide successfully models a premium Flask chatbot interface capable of answering basic college student "
        "and parent inquiries. By merging server-side python control with beautiful frontend UX elements and Gemini LLM "
        "integration, the bot delivers a highly robust interactive experience. Future enhancements will involve linking with "
        "live student databases via SQL libraries, supporting multi-language speech outputs, and fine-tuning models to "
        "automatically answer campus-specific curriculum documents."
    )

    doc.save(DOCX_OUT)
    print(f"Word file saved successfully to {DOCX_OUT}")


# --------------------------------------------------------------------------
# PDF (.pdf) Generation using ReportLab
# --------------------------------------------------------------------------
class NumberedCanvas(canvas.Canvas):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
        self._saved_page_states = []

    def showPage(self):
        self._saved_page_states.append(dict(self.__dict__))
        self._startPage()

    def save(self):
        num_pages = len(self._saved_page_states)
        for state in self._saved_page_states:
            self.__dict__.update(state)
            self.draw_page_decorations(num_pages)
            super().showPage()
        super().save()

    def draw_page_decorations(self, page_count):
        self.saveState()
        self.setFont("Helvetica", 9)
        self.setFillColor(colors.HexColor("#64748B"))
        
        if self._pageNumber > 1:
            # Header
            self.drawString(54, 750, "UniGuide - Flask AI Campus Assistant Chatbot")
            self.setStrokeColor(colors.HexColor("#E2E8F0"))
            self.setLineWidth(0.5)
            self.line(54, 742, 558, 742)
            
            # Footer
            self.line(54, 55, 558, 55)
            page_text = f"Page {self._pageNumber} of {page_count}"
            self.drawRightString(558, 40, page_text)
            self.drawString(54, 40, "Assignment 1: Chatbot Development - Flask AI Submission")
            
        self.restoreState()


def generate_pdf():
    print("Generating PDF file...")
    doc = SimpleDocTemplate(
        PDF_OUT,
        pagesize=letter,
        leftMargin=54,
        rightMargin=54,
        topMargin=72,
        bottomMargin=72
    )

    styles = getSampleStyleSheet()
    
    # Custom Paragraph Styles
    title_style = ParagraphStyle(
        'DocTitle',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=24,
        leading=28,
        textColor=colors.HexColor("#4F46E5"), # Indigo 600
        alignment=1, # Center
        spaceAfter=15
    )

    subtitle_style = ParagraphStyle(
        'DocSubtitle',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=14,
        leading=18,
        textColor=colors.HexColor("#64748B"), # Slate 500
        alignment=1, # Center
        spaceAfter=40
    )

    author_style = ParagraphStyle(
        'DocAuthor',
        parent=styles['Normal'],
        fontName='Helvetica-Bold',
        fontSize=11,
        leading=15,
        textColor=colors.HexColor("#0F172A"),
        alignment=1, # Center
        spaceAfter=5
    )

    h1_style = ParagraphStyle(
        'H1',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=16,
        leading=20,
        textColor=colors.HexColor("#4F46E5"),
        spaceBefore=18,
        spaceAfter=8,
        keepWithNext=True
    )

    h2_style = ParagraphStyle(
        'H2',
        parent=styles['Heading2'],
        fontName='Helvetica-Bold',
        fontSize=12,
        leading=16,
        textColor=colors.HexColor("#0F172A"),
        spaceBefore=14,
        spaceAfter=6,
        keepWithNext=True
    )

    body_style = ParagraphStyle(
        'Body',
        parent=styles['Normal'],
        fontName='Helvetica',
        fontSize=10,
        leading=14,
        textColor=colors.HexColor("#334155"),
        spaceAfter=8
    )

    bullet_style = ParagraphStyle(
        'Bullet',
        parent=body_style,
        leftIndent=15,
        firstLineIndent=-10,
        spaceAfter=4
    )

    caption_style = ParagraphStyle(
        'Caption',
        parent=styles['Normal'],
        fontName='Helvetica-Oblique',
        fontSize=9,
        leading=12,
        textColor=colors.HexColor("#64748B"),
        alignment=1, # Center
        spaceAfter=12
    )

    story = []

    # --- COVER PAGE ---
    story.append(Spacer(1, 150))
    story.append(Paragraph(title_text, title_style))
    story.append(Paragraph(subtitle_text, subtitle_style))
    story.append(Spacer(1, 100))
    story.append(Paragraph(author_text, author_style))
    story.append(Paragraph(f"Date: {date_text}", author_style))
    story.append(Paragraph("Architecture: Python Flask Backend & Gemini AI integration", author_style))
    story.append(PageBreak())

    # --- SECTION 1 ---
    story.append(Paragraph("1. Executive Summary & Project Objective", h1_style))
    story.append(Paragraph(
        "Chatbots have emerged as an essential interface in modern web systems, facilitating instant support "
        "and reducing administrative workload. This project, titled UniGuide, is a smart campus assistant chatbot "
        "built with a Python Flask server backend and optional Google Gemini AI integration. The chatbot is designed to guide "
        "students, prospective applicants, and parents through admissions procedures, tuition structures, academic "
        "calendars, campus geography, and support tickets.",
        body_style
    ))
    story.append(Paragraph("Key objectives of the UniGuide chatbot include:", body_style))
    story.append(Paragraph("• <b>24/7 Information Access:</b> Instantly answering common student questions regarding tuition fees and term dates.", bullet_style))
    story.append(Paragraph("• <b>Dual AI Processing:</b> Leveraging Google Gemini LLM API for natural language understanding while using a robust offline rule-based regex fallback.", bullet_style))
    story.append(Paragraph("• <b>Admissions Streamlining:</b> Helping new applicants understand application steps and deadlines.", bullet_style))
    story.append(Paragraph("• <b>Live Support Integration:</b> Enabling users to raise troubleshooting tickets, gathering email data, and generating confirmation IDs via stateful sessions.", bullet_style))

    # --- SECTION 2 ---
    story.append(Paragraph("2. Conversation Flow Design", h1_style))
    story.append(Paragraph(
        "To deliver a robust conversational user experience, UniGuide uses a rule-based state-machine design. "
        "The logic splits interactions into predefined, structured flows: Welcome Greeting, Admission Flow, Navigation Flow, "
        "Ticketing Form Flow, and Fallback/Error state.",
        body_style
    ))
    
    story.append(Paragraph("Detailed Conversation Node Structure", h2_style))
    story.append(Paragraph("1. <b>Welcome Node:</b> Greets the user, details the bot's abilities, and suggests quick-reply buttons (Admissions, Fees, Map, Exams, Raise Ticket).", bullet_style))
    story.append(Paragraph("2. <b>Admissions Node:</b> Triggers on keywords like 'apply', 'admission', 'enrol'. Displays steps, deadlines, and links to the Application portal.", bullet_style))
    story.append(Paragraph("3. <b>Navigation Node:</b> Triggers on map-related keywords. Displays a simulated graphical map and popular landmarks.", bullet_style))
    story.append(Paragraph("4. <b>Support Ticket State:</b> Collects student email, validates formatting, prompts for issue details, submits to HELP-DESK queue, and returns confirmation IDs.", bullet_style))
    story.append(Paragraph("5. <b>Fallback Error Node:</b> Triggers on unmatched queries. Suggests search alternatives and displays hotkeys to return to known topics.", bullet_style))

    # --- SECTION 3 ---
    story.append(Paragraph("3. Technical Implementation & UI Design", h1_style))
    story.append(Paragraph(
        "UniGuide is built as a highly responsive Python Flask application to guarantee server-side security, "
        "live API connectivity, and a premium look-and-feel. It uses native web APIs and complies with modern aesthetic guidelines.",
        body_style
    ))
    story.append(Paragraph("• <b>Backend Architecture (Python Flask):</b> Configures session state memory to store support ticket fields. Exposes routing for landing pages ('/') and API chat processing endpoints ('/api/chat'). Loads the Google Generative AI (Gemini) SDK for intelligent reasoning.", bullet_style))
    story.append(Paragraph("• <b>Structure (HTML5/Templates):</b> Rendered dynamically by Flask. Organizes the application into a split layout, showing stats cards in the sidebar and chat messages on the main viewport.", bullet_style))
    story.append(Paragraph("• <b>Design & Style (Vanilla CSS3):</b> Uses glassmorphism (frosted borders, backdrop filters) to blend components nicely in dark and light modes. Incorporates smooth transit keyframes for message bubble entries.", bullet_style))
    story.append(Paragraph("• <b>Client Script (Vanilla JS):</b> Manages frontend elements, triggers typing status indicators, interacts with the backend `/api/chat` API using Fetch, and plays sound effects.", bullet_style))

    story.append(PageBreak())

    # --- SECTION 4 ---
    story.append(Paragraph("4. Chatbot Screens & Output Analysis", h1_style))
    
    # Adding screenshots
    images_to_add = [
        ("welcome_screen.png", "Figure 1: UniGuide Welcome Interface & Common Topics Panel"),
        ("admission_faq.png", "Figure 2: Admissions Requirements & Fees Inquiry Interface"),
        ("navigation_flow.png", "Figure 3: Interactive Campus Maps and Directions"),
        ("fallback_handling.png", "Figure 4: Robust Fallback Error Handling for Unrecognized Queries")
    ]

    for filename, caption in images_to_add:
        img_path = images_ok.get(filename)
        if img_path and os.path.exists(img_path):
            img = Image(img_path, width=380, height=230)
            story.append(KeepTogether([img, Spacer(1, 4), Paragraph(caption, caption_style), Spacer(1, 10)]))
        else:
            story.append(Paragraph(f"[Image Missing: {filename}]", body_style))

    story.append(PageBreak())
    
    # Project Links
    story.append(Paragraph("Project Working Directory & Links", h2_style))
    project_dir_url = PROJECT_DIR.replace('\\', '/')
    flask_script_url = os.path.join(PROJECT_DIR, 'app.py').replace('\\', '/')
    story.append(Paragraph(f"• <b>Local Source Directory:</b> <code>file:///{project_dir_url}</code>", bullet_style))
    story.append(Paragraph(f"• <b>Flask Entry Script:</b> <code>file:///{flask_script_url}</code>", bullet_style))
    story.append(Paragraph("• <b>Local Dev Server Port:</b> <code>http://localhost:5000</code>", bullet_style))
    story.append(Spacer(1, 10))

    # Table of Dialog Flows
    story.append(Paragraph("Conversation Interaction Mapping", h2_style))
    
    data = [
        [
            Paragraph("<b>User Intent / Input</b>", body_style),
            Paragraph("<b>Pattern Matching Rule</b>", body_style),
            Paragraph("<b>Bot Action & Output</b>", body_style)
        ],
        [
            Paragraph("admission requirements", body_style),
            Paragraph("Keyword matches: 'admission', 'apply'", body_style),
            Paragraph("Displays application portal link, deadlines, and criteria list.", body_style)
        ],
        [
            Paragraph("tuition fees", body_style),
            Paragraph("Keyword matches: 'fees', 'cost', 'tuition'", body_style),
            Paragraph("Displays cost comparison list (in-state vs out-of-state) and scholarship links.", body_style)
        ],
        [
            Paragraph("campus map library", body_style),
            Paragraph("Keyword matches: 'map', 'directions', 'library'", body_style),
            Paragraph("Draws simulated campus map box with dynamic locator pulsing.", body_style)
        ],
        [
            Paragraph("raise ticket", body_style),
            Paragraph("Keyword matches: 'ticket', 'support', 'help'", body_style),
            Paragraph("Initiates stateful ticket form. Prompts for student email address, then prompt for issue description, returning ID.", body_style)
        ],
        [
            Paragraph("fly to Mars", body_style),
            Paragraph("No matches (Fallback triggers)", body_style),
            Paragraph("Returns friendly error notification card containing popular recommended buttons.", body_style)
        ]
    ]

    col_widths = [110, 140, 250]
    t = Table(data, colWidths=col_widths)
    t.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.HexColor("#E2E8F0")),
        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor("#CBD5E1")),
        ('PADDING', (0,0), (-1,-1), 8),
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
        ('BOTTOMPADDING', (0,0), (-1,0), 10),
    ]))
    story.append(t)
    story.append(Spacer(1, 20))

    # --- SECTION 5 ---
    story.append(Paragraph("5. Conclusion & Future Scopes", h1_style))
    story.append(Paragraph(
        "UniGuide successfully models a premium Flask chatbot interface capable of answering basic college student "
        "and parent inquiries. By merging server-side python control with beautiful frontend UX elements and Gemini LLM "
        "integration, the bot delivers a highly robust interactive experience. Future enhancements will involve linking with "
        "live student databases via SQL libraries, supporting multi-language speech outputs, and fine-tuning models to "
        "automatically answer campus-specific curriculum documents.",
        body_style
    ))

    # Build the document
    doc.build(story, canvasmaker=NumberedCanvas)
    print(f"PDF saved successfully to {PDF_OUT}")


if __name__ == "__main__":
    generate_docx()
    generate_pdf()
    print("Done! Both documents compiled.")
